from __future__ import annotations

import os
from dataclasses import dataclass, field
from typing import List, Optional

from docx import Document


# -------------------------
# Data models
# -------------------------
@dataclass
class ExperienceItem:
    title: str
    company: str
    start_date: str = ""
    end_date: str = ""
    location: str = ""
    bullets: List[str] = field(default_factory=list)


@dataclass
class EducationItem:
    degree: str
    institution: str
    year: str = ""
    details: str = ""


@dataclass
class ResumeData:
    full_name: str
    headline: str
    location: str
    email: str
    phone: str
    links: List[str]
    summary: str
    skills: List[str]
    experience: List[ExperienceItem]
    education: List[EducationItem]
    certifications: List[str] = field(default_factory=list)
    projects: List[str] = field(default_factory=list)


# -------------------------
# Renderer
# -------------------------
class ResumeTemplateRenderer:
    """
    Minimal DOCX renderer aligned with your app.py expectations:
    - Accepts output_dir in __init__
    - Provides render_modern / render_executive / render_creative / render_technical
    - Each returns a file path to the generated .docx
    """

    def __init__(self, output_dir: str = "generated_documents") -> None:
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def _build_doc(self, data: ResumeData, style_name: str) -> Document:
        """
        Internal builder. 'style_name' is just a label here (modern/executive/etc).
        You can later customize formatting per template.
        """
        doc = Document()

        # Header
        doc.add_heading(data.full_name, level=0)
        if data.headline:
            doc.add_paragraph(data.headline)

        contact_line = " | ".join(
            [x for x in [data.location, data.email, data.phone] if x]
        )
        if contact_line:
            doc.add_paragraph(contact_line)

        if data.links:
            doc.add_paragraph("Links: " + " | ".join(data.links))

        doc.add_paragraph(f"Template: {style_name}")

        # Summary
        if data.summary:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(data.summary)

        # Skills
        if data.skills:
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(data.skills))

        # Experience
        if data.experience:
            doc.add_heading("Experience", level=1)
            for ex in data.experience:
                line = f"{ex.title} — {ex.company}"
                dates = " – ".join([d for d in [ex.start_date, ex.end_date] if d])
                if dates:
                    line += f" ({dates})"
                if ex.location:
                    line += f" | {ex.location}"

                doc.add_paragraph(line)

                for b in ex.bullets:
                    p = doc.add_paragraph(b)
                    p.style = "List Bullet"

        # Education
        if data.education:
            doc.add_heading("Education", level=1)
            for ed in data.education:
                line = f"{ed.degree} — {ed.institution}"
                if ed.year:
                    line += f" ({ed.year})"
                doc.add_paragraph(line)
                if ed.details:
                    doc.add_paragraph(ed.details)

        # Certifications
        if data.certifications:
            doc.add_heading("Certifications", level=1)
            for c in data.certifications:
                p = doc.add_paragraph(c)
                p.style = "List Bullet"

        # Projects
        if data.projects:
            doc.add_heading("Projects", level=1)
            for pr in data.projects:
                p = doc.add_paragraph(pr)
                p.style = "List Bullet"

        return doc

    def _save(self, doc: Document, filename: str) -> str:
        out_path = os.path.join(self.output_dir, filename)
        doc.save(out_path)
        return out_path

    # Public API expected by app.py
    def render_modern(self, data: ResumeData, filename: str = "resume_modern.docx") -> str:
        doc = self._build_doc(data, style_name="modern")
        return self._save(doc, filename)

    def render_executive(self, data: ResumeData, filename: str = "resume_executive.docx") -> str:
        doc = self._build_doc(data, style_name="executive")
        return self._save(doc, filename)

    def render_creative(self, data: ResumeData, filename: str = "resume_creative.docx") -> str:
        doc = self._build_doc(data, style_name="creative")
        return self._save(doc, filename)

    def render_technical(self, data: ResumeData, filename: str = "resume_technical.docx") -> str:
        doc = self._build_doc(data, style_name="technical")
        return self._save(doc, filename)


# Optional: keep a simple text renderer if you still want it
def render_resume(data: dict) -> str:
    name = data.get("name", "Unknown")
    summary = data.get("summary", "")
    title = data.get("title", "")
    return f"{name}\n{title}\n\nSUMMARY\n{summary}\n"
