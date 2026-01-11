from __future__ import annotations

import os
from dataclasses import dataclass, field
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


@dataclass
class ExperienceItem:
    title: str
    company: str
    start_date: str
    end_date: str
    location: str = ""
    bullets: List[str] = field(default_factory=list)


@dataclass
class EducationItem:
    degree: str
    institution: str
    year: str = ""
    details: str = ""


@dataclass
class ResumeData:
    full_name: str
    headline: str
    location: str
    email: str
    phone: str
    links: List[str]
    summary: str
    skills: List[str]
    experience: List[ExperienceItem]
    education: List[EducationItem]
    certifications: List[str] = field(default_factory=list)
    projects: List[str] = field(default_factory=list)


class ResumeTemplateRenderer:
    def __init__(self, output_dir: str = "generated_documents", brand_line: Optional[str] = None):
        self.output_dir = output_dir
        self.brand_line = brand_line
        os.makedirs(self.output_dir, exist_ok=True)

    def _add_footer_branding(self, doc: Document) -> None:
        if not self.brand_line:
            return
        for section in doc.sections:
            footer = section.footer
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            if p.text.strip():
                p = footer.add_paragraph()
            p.text = self.brand_line
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].font.size = Pt(8)

    # --- Example renderers (keep yours if you already have real templates) ---
    def render_modern(self, data: ResumeData, filename: str) -> str:
        path = os.path.join(self.output_dir, filename)
        doc = Document()
        doc.add_heading(data.full_name, level=0)
        doc.add_paragraph(data.headline)
        doc.add_paragraph(f"{data.location} • {data.email} • {data.phone}")

        if data.links:
            doc.add_paragraph(" | ".join(data.links))

        doc.add_heading("Summary", level=1)
        doc.add_paragraph(data.summary)

        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(data.skills))

        doc.add_heading("Experience", level=1)
        for ex in data.experience:
            doc.add_paragraph(f"{ex.title} — {ex.company} ({ex.start_date}–{ex.end_date})")
            if ex.location:
                doc.add_paragraph(ex.location)
            for b in ex.bullets:
                doc.add_paragraph(f"• {b}")

        doc.add_heading("Education", level=1)
        for ed in data.education:
            line = f"{ed.degree} — {ed.institution}"
            if ed.year:
                line += f" ({ed.year})"
            doc.add_paragraph(line)
            if ed.details:
                doc.add_paragraph(ed.details)

        if data.certifications:
            doc.add_heading("Certifications", level=1)
            for c in data.certifications:
                doc.add_paragraph(f"• {c}")

        if data.projects:
            doc.add_heading("Projects", level=1)
            for p in data.projects:
                doc.add_paragraph(f"• {p}")

        self._add_footer_branding(doc)
        doc.save(path)
        return path

    # If you have different templates, you can route them to render_modern for now:
    def render_executive(self, data: ResumeData, filename: str) -> str:
        return self.render_modern(data, filename)

    def render_creative(self, data: ResumeData, filename: str) -> str:
        return self.render_modern(data, filename)

    def render_technical(self, data: ResumeData, filename: str) -> str:
        return self.render_modern(data, filename)
