from __future__ import annotations

import os
from dataclasses import dataclass, field
from typing import List

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---- Data models ----
@dataclass
class ExperienceItem:
    title: str
    company: str
    start_date: str = ""
    end_date: str = ""
    location: str = ""
    bullets: List[str] = field(default_factory=list)

@dataclass
class EducationItem:
    degree: str
    institution: str
    year: str = ""
    details: str = ""

@dataclass
class ResumeData:
    full_name: str
    headline: str
    location: str
    email: str
    phone: str
    links: List[str]
    summary: str
    skills: List[str]
    experience: List[ExperienceItem]
    education: List[EducationItem]
    certifications: List[str] = field(default_factory=list)
    projects: List[str] = field(default_factory=list)


# ---- Renderer ----
class ResumeTemplateRenderer:
    def __init__(self, output_dir: str = "generated_documents", brand_line: str | None = None) -> None:
        self.output_dir = output_dir
        self.brand_line = brand_line
        os.makedirs(self.output_dir, exist_ok=True)

    def _add_brand_footer(self, doc: Document) -> None:
        if not self.brand_line:
            return
        for section in doc.sections:
            footer = section.footer
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.text = self.brand_line
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].font.size = Pt(8)

    def _build_doc(self, data: ResumeData, style_name: str) -> Document:
        doc = Document()
        doc.add_heading(data.full_name, level=0)
        if data.headline:
            doc.add_paragraph(data.headline)

        contact_line = " | ".join([x for x in [data.location, data.email, data.phone] if x])
        if contact_line:
            doc.add_paragraph(contact_line)

        if data.links:
            doc.add_paragraph("Links: " + " | ".join(data.links))

        if data.summary:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(data.summary)

        if data.skills:
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(data.skills))

        if data.experience:
            doc.add_heading("Experience", level=1)
            for ex in data.experience:
                line = f"{ex.title} — {ex.company}"
                dates = " – ".join([d for d in [ex.start_date, ex.end_date] if d])
                if dates:
                    line += f" ({dates})"
                if ex.location:
                    line += f" | {ex.location}"
                doc.add_paragraph(line)
                for b in ex.bullets:
                    p = doc.add_paragraph(b)
                    p.style = "List Bullet"

        if data.education:
            doc.add_heading("Education", level=1)
            for ed in data.education:
                line = f"{ed.degree} — {ed.institution}"
                if ed.year:
                    line += f" ({ed.year})"
                doc.add_paragraph(line)
                if ed.details:
                    doc.add_paragraph(ed.details)

        if data.certifications:
            doc.add_heading("Certifications", level=1)
            for c in data.certifications:
                p = doc.add_paragraph(c)
                p.style = "List Bullet"

        if data.projects:
            doc.add_heading("Projects", level=1)
            for pr in data.projects:
                p = doc.add_paragraph(pr)
                p.style = "List Bullet"

        return doc

    def _save(self, doc: Document, filename: str) -> str:
        self._add_brand_footer(doc)
        out_path = os.path.join(self.output_dir, filename)
        doc.save(out_path)
        return out_path

    def render_modern(self, data: ResumeData, filename: str) -> str:
        return self._save(self._build_doc(data, "modern"), filename)

    def render_executive(self, data: ResumeData, filename: str) -> str:
        return self._save(self._build_doc(data, "executive"), filename)

    def render_creative(self, data: ResumeData, filename: str) -> str:
        return self._save(self._build_doc(data, "creative"), filename)

    def render_technical(self, data: ResumeData, filename: str) -> str:
        return self._save(self._build_doc(data, "technical"), filename)
